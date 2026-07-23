from docx import Document
from pathlib import Path

path = Path('SMS_REMINDER_SYSTEM_updated.docx')
source = Document(path)
paragraphs = source.paragraphs

# Find the important chapter markers.
chapter4_idx = None
chapter5_idx = None
references_idx = None
for i, p in enumerate(paragraphs):
    text = p.text.strip().upper()
    if text == 'CHAPTER FOUR':
        chapter4_idx = i
    elif text == 'CHAPTER FIVE':
        chapter5_idx = i
    elif text == 'REFERENCES':
        references_idx = i
        break

if chapter4_idx is None or chapter5_idx is None or references_idx is None:
    raise SystemExit('Required chapter markers were not found in the document.')

# Keep everything before Chapter 4 and everything from References onward.
head = [p.text for p in paragraphs[:chapter4_idx]]
footer = [p.text for p in paragraphs[references_idx:]]

# New Chapter 4 and 5 body text.
chapter4 = [
    'CHAPTER FOUR',
    'SYSTEM IMPLEMENTATION AND TESTING',
    '4.0 Introduction',
    'This chapter presents how the Automated SMS Timetable Reminder System was implemented and tested in the course of the project. It explains the tools used in building the system, describes the main functional modules that were developed, and reports the testing activities that were carried out to confirm whether the solution was practical and useful. The development process followed the Agile approach, which allowed features to be built in short cycles and improved through feedback from the users.',
    '4.1 Implementation Environment',
    'The system was implemented using a combination of web technologies and cloud services that were suitable for building a secure and responsive application. Next.js was used for the frontend and backend structure, while TypeScript was used to support clearer and more reliable code. Tailwind CSS was used to style the administrative and lecturer interfaces in a simple and user-friendly way. Supabase was selected as the database platform because it provides PostgreSQL storage, authentication and row-level security in a single environment. Termii was used as the SMS gateway for sending reminders to lecturers. The application was hosted on Vercel, while GitHub was used for version control and Visual Studio Code served as the main code editor during development.',
    '4.2 Description of System Modules',
    'The system was delivered through six main functional modules. The first five were planned from the beginning, while the assistant module was added later after feedback showed that lecturers wanted a quicker way to ask direct questions about their timetable.',
    'Authentication Module. This module handles login and session management for both lecturers and administrators through Supabase Auth. Each user account is linked to a profile record that stores the user role and other basic information. The lecturer and administrator login pages were kept separate to maintain a clearer security boundary. Some administrators also teach classes, so a switch was added to let them move into the lecturer view without needing to log out and log back in.',
    'Timetable Management Module. This module gives the administrator the ability to create, edit and manage timetable entries. Each timetable record stores the course code, course name, assigned lecturer, day of the week, start time, end time and venue. Entries are not completely deleted when they are no longer active; instead, they are marked inactive so that past timetable records can still be preserved while the active reminder system continues to work properly.',
    'SMS Scheduling and Dispatch Module. This is the central module of the system. A scheduled background job checks the timetable regularly to identify classes that are due for reminder messages. The job compares the timetable with the current time using the correct time zone for Nigeria, so that reminder timing remains accurate. Before a message is sent, the system checks whether that same class has already received a reminder for the day in order to avoid duplicate messages. Once a reminder is due, the message is formatted and sent through the Termii API. The delivery outcome is recorded in the notifications table so that the administrator can see whether the reminder was delivered or failed.',
    'Lecturer Self-Service Module. This module gives each lecturer access to a personal timetable view that shows only the classes assigned to that lecturer. It also provides a change request form that allows a lecturer to select a class, give a reason for the request and suggest a new date where necessary. Submitted requests are sent to the administrator for review and approval or rejection.',
    'Notification Log Module. This module provides the administrator with a clear record of all SMS reminders that have been attempted. The log shows the recipient, message content, time sent and delivery status, making it easy to identify failed messages and correct them quickly. Lecturers can also view only the notifications related to their own account.',
    'Assistant Module. During testing, it became clear that both lecturers and administrators wanted a faster way to ask direct questions such as “How many classes do I have today?” instead of navigating through the dashboard to find the answer. The assistant was therefore added as a lightweight feature that matches user input against expected timetable-related patterns in the database. It does not rely on an external large language model. Instead, it uses the same data already stored in Supabase to answer simple questions quickly and at a low cost.',
    '4.3 Database Implementation',
    'The database was implemented in Supabase using four core tables that support the main functions of the system. The profiles table stores each user’s full name, phone number, email address, role and other basic information. The timetable table stores all class schedule records, including the course code, course name, lecturer, day, start time, end time and venue. The change_requests table stores lecturer requests for timetable changes, including the status of the request and any feedback from the administrator. The notifications table stores a permanent log of every SMS reminder attempt, including the time sent and whether the message was delivered or failed. Row-level security was enabled on all four tables so that lecturers can only access their own records, while administrators have access to the wider system when required. This provides a strong level of control over data access and helps protect the system from unauthorised viewing.',
    '4.4 System Testing',
    'Testing was carried out in stages during development and after the major modules had been integrated. This made it possible to detect problems early and correct them before the system was finalised.',
    '4.4.1 Unit Testing',
    'Each module was tested separately as it was completed. The login pages were verified to ensure that a lecturer account could not be used on the administrator page and vice versa. The row-level security setting was also checked to confirm that one lecturer could not view another lecturer’s timetable or change requests. The SMS message builder was tested to ensure that the course code, course name, start time and venue were correctly formatted into a readable reminder message. The reminder scheduling logic and assistant logic were also tested to confirm that the correct time zone was being used and that schedule-related questions were interpreted properly.',
    '4.4.2 Integration Testing',
    'Integration testing was used to confirm that the different modules worked together as one system. The flow began with an administrator creating or updating a timetable entry, followed by the scheduler identifying the entry when it was due, the Termii API sending the reminder to the lecturer’s phone number, the result being recorded in the notifications table, and the administrator and lecturer being able to view the outcome through their respective interfaces. This testing stage helped confirm that the system worked as a connected application rather than as separate standalone parts.',
    '4.4.3 User Acceptance Testing',
    'A questionnaire was used earlier during requirements gathering to understand how timetable information was being shared and what lecturers expected from a better reminder system. After the system had been built, it was demonstrated to four lecturers from the Computer Science Department. These lecturers were not required to sign in during the development trial because the goal at that stage was mainly to observe how the system behaved, how the schedule view worked, and whether the reminder process was practical and understandable from the lecturers’ point of view. They were asked to view their timetable, test the change request feature, and observe the reminder flow. Their feedback was used to improve the system, especially by refining the request form and adding the assistant feature after lecturers requested a quicker way to check their schedule.',
    '4.4.4 SMS Delivery Testing',
    'Because Termii may route messages differently depending on the phone network provider, test messages were sent to numbers on the major mobile networks in Nigeria, including MTN, Glo, Airtel and 9Mobile, to confirm that delivery worked across the different networks. Delivery was confirmed through the system’s notification log and the Termii dashboard, which reports the outcome of every message sent through the platform.',
    '4.5 Discussion of Test Results',
    'The testing process showed that the system met the main objectives stated in Chapter One. Automatic SMS reminders were successfully delivered to lecturers before their classes began, reducing the need for manual timetable distribution. The lecturer portal gave lecturers access to their own schedule and allowed them to submit change requests without visiting the department office. The administrator dashboard made it possible to monitor which reminders were delivered and which failed. The testing also revealed two important issues. The first was the difference between the server time zone and the local time zone in Nigeria, which could affect reminder timing if the system was not adjusted correctly. The second concerned the assistant feature, where some natural phrases were not recognised properly. Both issues were corrected during testing, which showed the importance of dealing with real user behaviour instead of relying only on design assumptions.',
]

chapter5 = [
    'CHAPTER FIVE',
    'SUMMARY, CONCLUSION AND RECOMMENDATIONS',
    '5.0 Introduction',
    'This final chapter summarises the study, presents the main conclusions drawn from the design, implementation and testing of the Automated SMS Timetable Reminder System, and offers recommendations for improving the system and extending it to other departments if needed.',
    '5.1 Summary of the Study',
    'This study was carried out to address the problem of manual timetable distribution in the Computer Science Department at the National Institute of Construction Technology and Management, Uromi. Lecturers often depended on printed timetables and verbal reminders, which made it easy for important information to be missed or forgotten. The project therefore focused on developing a system that would automatically send reminder messages to lecturers before their classes began, while also giving them a simple way to view their timetable and submit change requests. The Agile development model was adopted because it allowed the system to be built in cycles, with testing and feedback happening continuously throughout the project.',
    '5.2 Summary of Findings',
    'The completed system achieved the three main objectives stated in Chapter One. First, it automated the sending of SMS reminders to lecturers ahead of class time. Second, it allowed lecturers to view their personal timetable and submit change requests through a self-service portal. Third, it gave administrators a clear notification log showing which reminders were delivered and which failed. During the development process, the assistant feature was also added after lecturers requested a faster method for finding schedule-related answers. This showed that the system was responsive to real user need rather than being limited only to the original plan.',
    '5.3 Limitations Encountered During Implementation',
    'The study also identified some limitations that affected the full implementation of the system. The system depends on Termii as a third-party SMS provider, so any downtime or delay on the service side directly affects delivery. There is no dedicated mobile application; lecturers access the portal through a mobile browser. The system was implemented and tested within the Computer Science Department only, so its performance in other departments has not yet been evaluated. In addition, reminder delivery depends on the accuracy of each lecturer’s registered phone number. If a number is outdated or incorrect, the reminder may fail even though the system has recorded the attempt.',
    '5.4 Conclusion',
    'This project has shown that a low-cost, cloud-based SMS reminder system can make timetable communication more reliable and more efficient for lecturers. By combining a web application, a structured database and an SMS gateway, the system reduces the manual effort involved in distributing and updating timetable information while giving administrators better oversight of delivery results. The project also shows that useful features do not always require complex technologies. The assistant module added value by answering simple questions quickly using the existing database instead of a costly external AI service.',
    '5.5 Recommendations',
    'The system should be extended to other departments at NICTM so that the model can be evaluated on a wider scale. A dedicated mobile application should also be considered in the future to make the lecturer experience even more convenient. The department should continue to ensure that lecturer phone numbers are updated regularly, since this is one of the most common reasons for failed reminders. Finally, if the system is to be used on a larger scale, the department should consider moving to a paid Termii plan so that message volume can be handled more effectively.',
    '5.6 Suggestions for Further Studies',
    'Future research could examine a two-way SMS feature that allows lecturers to reply to reminder messages to confirm that they have seen them. Future work could also focus on data analytics that show delivery patterns and request trends over a whole academic session. Another useful area for further study would be the use of a small locally hosted language model to improve the assistant’s ability to handle more varied user questions while still keeping the system affordable and secure.',
]

# Build a fresh document to avoid broken paragraph references.
new_doc = Document()

# Recreate all content before Chapter 4.
for line in head:
    new_doc.add_paragraph(line)

# Add the revised Chapter 4 and Chapter 5.
for line in chapter4 + chapter5:
    new_doc.add_paragraph(line)

# Add the references section from the original file.
for line in footer:
    new_doc.add_paragraph(line)

new_doc.save(path)
print('Chapter 4 and 5 text updated successfully.')
